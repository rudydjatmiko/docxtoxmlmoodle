from docx import Document
from lxml import etree
import re


# =========================
# NORMALISASI SIMBOL
# =========================
def normalize_symbol(text):
    """
    Pastikan simbol UTF-8 aman
    """
    if not text:
        return ""
    return text.encode("utf-8", "ignore").decode()


# =========================
# CLEAN TEXT UMUM
# =========================
def clean_text(text):
    """
    Bersihkan text aneh dari Word
    """
    if not text:
        return ""

    text = normalize_symbol(text)

    # hapus karakter aneh
    text = re.sub(r'[\u200b-\u200f]', '', text)

    return text.strip()


# =========================
# EXTRACT TABLE → HTML
# =========================
def extract_tables(docx_file):
    """
    Ambil semua tabel dari docx → HTML
    """
    docx_file.seek(0)
    doc = Document(docx_file)

    tables_html = []

    for table in doc.tables:

        html = "<table border='1' style='border-collapse:collapse;'>"

        for row in table.rows:
            html += "<tr>"
            for cell in row.cells:
                cell_text = clean_text(cell.text)
                html += f"<td>{cell_text}</td>"
            html += "</tr>"

        html += "</table>"

        tables_html.append(html)

    return tables_html


# =========================
# DETEKSI TABLE DI TEXT
# =========================
def inject_table_to_text(text, tables):
    """
    Sisipkan tabel ke text jika ditemukan marker
    """
    if not tables:
        return text

    # contoh marker (opsional)
    for i, table_html in enumerate(tables, start=1):
        marker = f"[TABLE{i}]"
        if marker in text:
            text = text.replace(marker, table_html)

    return text


# =========================
# EXTRACT TEXTBOX (ADVANCED)
# =========================
def extract_textbox(docx_file):
    """
    Ambil text dari textbox (shape)
    """
    docx_file.seek(0)
    doc = Document(docx_file)

    xml = doc._element.xml
    tree = etree.fromstring(xml.encode())

    texts = tree.xpath('//w:t')

    results = []
    for t in texts:
        if t.text:
            cleaned = clean_text(t.text)
            if cleaned:
                results.append(cleaned)

    return results


# =========================
# PROCESS FINAL CONTENT
# =========================
def process_content(text, docx_file=None):
    """
    Pipeline konten:
    - clean text
    - inject tabel (jika ada)
    """

    text = clean_text(text)

    if docx_file:
        tables = extract_tables(docx_file)
        text = inject_table_to_text(text, tables)

    return text
