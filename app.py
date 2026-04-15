import sys
import os

# ======================
# FIX PATH (WAJIB)
# ======================
BASE_DIR = os.path.dirname(__file__)
sys.path.insert(0, BASE_DIR)

# ======================
# IMPORT
# ======================
import streamlit as st
from docx import Document
from core.docx_reader import read_docx
from parser import run_parser
from core.builder import build_xml
from utils.xml_parser import get_xml_info

# ======================
# CONFIG
# ======================
st.set_page_config(
    page_title="DOCX to Moodle XML",
    page_icon="📄",
    layout="centered"
)

# ======================
# TITLE
# ======================
st.title("📄 DOCX → Moodle XML Converter")
st.caption("Convert DOCX ke XML Moodle + Debug Tools (RAW, Baris, XML)")

# ======================
# SESSION STATE
# ======================
if "xml_result" not in st.session_state:
    st.session_state.xml_result = None

# ======================
# UPLOAD
# ======================
st.markdown("### 📥 Upload File DOCX")

uploaded_file = st.file_uploader(
    "Upload file",
    type=["docx"],
    label_visibility="collapsed"
)

# ======================
# HELPER
# ======================
def save_temp(uploaded_file):
    temp_path = os.path.join(BASE_DIR, "temp.docx")
    with open(temp_path, "wb") as f:
        f.write(uploaded_file.read())
    return temp_path

# ======================
# BUTTONS
# ======================
col1, col2, col3, col4, col5 = st.columns(5)

process_btn = col1.button("🚀 Proses")
reset_btn = col2.button("🔄 Reset")
debug_raw_btn = col3.button("📄 RAW")
debug_lines_btn = col4.button("🔍 Baris")
debug_xml_btn = col5.button("🧬 XML")

# ======================
# RESET
# ======================
if reset_btn:
    st.session_state.xml_result = None

    temp_path = os.path.join(BASE_DIR, "temp.docx")
    if os.path.exists(temp_path):
        os.remove(temp_path)

    st.rerun()

# ======================
# DEBUG RAW (ASLI DOCX)
# ======================
if uploaded_file and debug_raw_btn:

    with st.spinner("🔍 Membaca RAW DOCX..."):

        try:
            temp_path = save_temp(uploaded_file)
            doc = Document(temp_path)

            st.subheader("📄 RAW TEXT (ASLI)")

            lines = []
            for para in doc.paragraphs:
                text = para.text if para.text else "[KOSONG]"
                lines.append(text)

            st.code("\n".join(lines))

        except Exception as e:
            st.error(f"❌ Error RAW: {str(e)}")

# ======================
# DEBUG BARIS
# ======================
if uploaded_file and debug_lines_btn:

    with st.spinner("🔍 Membaca Baris..."):

        try:
            temp_path = save_temp(uploaded_file)
            elements = read_docx(temp_path)

            st.subheader("🔍 DEBUG BARIS")

            lines = []

            for i, el in enumerate(elements):
                text = el.get("text", "").strip()
                if not text:
                    text = "[KOSONG]"

                img_count = len(el.get("images", []))

                line = f"{i+1:03d} | {text}"

                if img_count > 0:
                    line += f"  🖼️({img_count})"

                if "ANS:" in text:
                    line = "🔑 " + line

                elif len(text.split()) <= 5:
                    line = "👉 " + line

                lines.append(line)

            st.code("\n".join(lines))

        except Exception as e:
            st.error(f"❌ Error Baris: {str(e)}")

# ======================
# DEBUG XML (NUMPR + DRAWING)
# ======================
if uploaded_file and debug_xml_btn:

    with st.spinner("🧬 Membaca XML DOCX..."):

        try:
            temp_path = save_temp(uploaded_file)
            doc = Document(temp_path)

            st.subheader("🧬 DEBUG XML (w:numPr & w:drawing)")

            output = []

            for i, para in enumerate(doc.paragraphs):

                text = para.text.strip() if para.text else "[KOSONG]"

                numbering, has_drawing = get_xml_info(para)

                line = f"{i+1:03d} | {text}"

                if numbering:
                    line += f"\n     numPr: level={numbering['level']} numId={numbering['numId']}"

                if has_drawing:
                    line += "\n     🖼️ drawing detected"

                output.append(line)

            st.code("\n\n".join(output))

        except Exception as e:
            st.error(f"❌ Error XML: {str(e)}")

# ======================
# PROCESS DOCX → XML
# ======================
if uploaded_file and process_btn:

    with st.spinner("🔄 Memproses DOCX → XML..."):

        try:
            temp_path = save_temp(uploaded_file)

            elements = read_docx(temp_path)
            questions = run_parser(elements)
            xml = build_xml(questions)

            st.session_state.xml_result = xml

        except Exception as e:
            st.error(f"❌ Error: {str(e)}")

# ======================
# RESULT
# ======================
if st.session_state.xml_result:

    st.success("✅ Konversi berhasil!")

    original_name = uploaded_file.name
    base_name = os.path.splitext(original_name)[0]
    xml_filename = base_name + ".xml"

    st.download_button(
        label="⬇️ Download XML",
        data=st.session_state.xml_result,
        file_name=xml_filename,
        mime="text/xml",
        use_container_width=True
    )

    with st.expander("🔍 Preview XML"):
        st.code(st.session_state.xml_result[:2000], language="xml")
